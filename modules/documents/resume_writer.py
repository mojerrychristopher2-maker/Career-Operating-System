from pathlib import Path
from docx import Document


class ResumeWriter:

    def __init__(self):
        self.output_dir = Path("output/resumes")
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def create(self, resume):

        document = Document()

        document.add_heading(resume["candidate"], level=1)

        document.add_paragraph(resume["headline"])

        document.add_paragraph(resume["location"])

        document.add_heading("Professional Summary", level=2)
        document.add_paragraph(resume["summary"])

        document.add_heading("Skills", level=2)

        for skill in resume["skills"]:
            document.add_paragraph(skill, style="List Bullet")

        document.add_heading("Experience", level=2)

        for exp in resume["experience"]:
            document.add_paragraph(exp, style="List Bullet")

        document.add_heading("Education", level=2)

        for edu in resume["education"]:
            document.add_paragraph(edu)

        document.add_heading("Certifications", level=2)

        for cert in resume["certifications"]:
            document.add_paragraph(cert, style="List Bullet")

        import re

        company = resume.get("company", "UnknownCompany")
        title = resume.get("job_title", "UnknownRole")

        safe_company = re.sub(r'[\\/:*?"<>|]', "", company).replace(" ", "_")
        safe_title = re.sub(r'[\\/:*?"<>|]', "", title).replace(" ", "_")

        filename = f"{safe_company}_{safe_title}.docx"

        filepath = self.output_dir / filename

        document.save(filepath)

        return filepath