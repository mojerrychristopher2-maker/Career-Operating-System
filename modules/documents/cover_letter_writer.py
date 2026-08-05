from pathlib import Path
from docx import Document


class CoverLetterWriter:

    def __init__(self):
        self.output_dir = Path("output/cover_letters")
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def create(self, cover_letter):

        document = Document()

        document.add_heading("Cover Letter", level=1)

        document.add_paragraph(cover_letter)

        filename = "Cover_Letter.docx"

        filepath = self.output_dir / filename

        document.save(filepath)

        return filepath